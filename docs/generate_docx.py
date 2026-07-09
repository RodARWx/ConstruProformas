import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc_path = r"c:\Users\gt\Documents\Sexto Semestre-3 de Septimo - 06-04-2026\Sexto\Desarrollo De Sistemas de Informacion\construmetrica creacion de programa sobre la preformas\ContruProformas\docs\REQUISITOS_SISTEMA.md"
    docx_path = r"c:\Users\gt\Documents\Sexto Semestre-3 de Septimo - 06-04-2026\Sexto\Desarrollo De Sistemas de Informacion\construmetrica creacion de programa sobre la preformas\ContruProformas\docs\REQUISITOS_SISTEMA.docx"
    
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found")
        return
        
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    with open(doc_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line_str = line.strip('\n')
        line_stripped = line_str.strip()
        
        if not line_stripped:
            continue
            
        # Headers
        if line_stripped.startswith('# '):
            title_text = line_stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00) # Dark red / Construmétrica color
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
        elif line_stripped.startswith('## '):
            h_text = line_stripped[3:]
            p = doc.add_heading(level=1)
            run = p.runs[0] if p.runs else p.add_run(h_text)
            run.text = h_text
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A) # Blue color
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            
        elif line_stripped.startswith('### '):
            h_text = line_stripped[4:]
            p = doc.add_heading(level=2)
            run = p.runs[0] if p.runs else p.add_run(h_text)
            run.text = h_text
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            
        elif line_stripped.startswith('---'):
            continue
            
        elif line_stripped.startswith('* ') or line_stripped.startswith('- '):
            bullet_text = line_stripped[2:]
            is_sub = line_str.startswith('  * ') or line_str.startswith('  - ') or line_str.startswith('\t* ') or line_str.startswith('\t- ')
            
            p = doc.add_paragraph(style='List Bullet 2' if is_sub else 'List Bullet')
            p.paragraph_format.space_after = Pt(2)
            
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parts = re.split(r'(\*\*.*?\*\*)', line_stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(docx_path)
    print(f"Word document saved to {docx_path}")

if __name__ == '__main__':
    main()
